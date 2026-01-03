import uuid
from fastapi import UploadFile
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from app.models.document import Document
from app.services.schema import get_chroma_collection, add_document_to_collection
from sentence_transformers import SentenceTransformer

class IngestionService:
    def __init__(self):
        """Initialize the IngestionService with ChromaDB collection and embedder."""
        self.collection = get_chroma_collection("documents")
        self.embedder = SentenceTransformer("all-MiniLM-L6-v2")

    def ingest_file(self, file: UploadFile):
        """Ingest a PDF or DOCX file, generate embedding, and store in ChromaDB."""
        filename = file.filename
        filetype = filename.split('.')[-1].lower()
        content = ""
        title = filename

        if filetype == "pdf":
            content = self._parse_pdf(file)
        elif filetype == "docx":
            content = self._parse_docx(file)
        else:
            raise ValueError("Unsupported file type")

        doc_id = str(uuid.uuid4())
        embedding = self.embedder.encode([content])[0]
        metadata = {"id": doc_id, "title": title, "filename": filename, "filetype": filetype}
        add_document_to_collection(self.collection, embedding, content, metadata, doc_id)
        print(self.collection.count())
        print(self.collection.get(include=["documents", "metadatas"]))
        document = Document(
            id=doc_id,
            title=title,
            content=content,
            metadata={"filename": filename, "filetype": filetype}
        )
        return document

    def _parse_pdf(self, file: UploadFile) -> str:
        """Extract text content from a PDF file."""
        file.file.seek(0)
        reader = PdfReader(file.file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def _parse_docx(self, file: UploadFile) -> str:
        """Extract text content from a DOCX file."""
        file.file.seek(0)
        doc = DocxDocument(file.file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text

    def list_documents(self):
        results = self.collection.get(include=["metadatas", "documents"])
        print(self.collection.count())
        documents = []
        for i, doc_id in enumerate(results.get("ids", [])):
            meta = results["metadatas"][i] if results["metadatas"] else {}
            content = results["documents"][i] if results["documents"] else None

            documents.append(Document(
                id=doc_id,
                title=meta.get("title", meta.get("filename", "")),
                content=content,
                metadata={
                    "filename": meta.get("filename", ""),
                    "filetype": meta.get("filetype", "")
                }
            ))
        return documents